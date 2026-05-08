"""
courses/services/document_parser.py
~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
Сервис извлечения текста из загружаемых преподавателем файлов.

Поддерживаемые форматы: .pdf, .docx
Зависимости: pypdf>=4.0.0 (уже в pyproject.toml), python-docx (добавить)
"""

from __future__ import annotations

import logging
from typing import BinaryIO

logger = logging.getLogger(__name__)

# Разрешённые расширения файлов
ALLOWED_EXTENSIONS = {".pdf", ".docx"}


# ---------------------------------------------------------------------------
# Исключения
# ---------------------------------------------------------------------------

class UnsupportedFileTypeError(ValueError):
    """Расширение файла не поддерживается сервисом."""


class ParsingError(RuntimeError):
    """Файл повреждён или не может быть прочитан."""


# ---------------------------------------------------------------------------
# Парсеры по типу файла
# ---------------------------------------------------------------------------

def _parse_pdf(file_obj: BinaryIO) -> str:
    """
    Извлекает текст из PDF-файла постранично.

    Использует pypdf.PdfReader. Пустые страницы пропускаются.
    При ошибке извлечения конкретной страницы — логируем, продолжаем.
    """
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError(
            "Библиотека pypdf не установлена. Выполните: poetry add pypdf"
        ) from exc

    try:
        reader = PdfReader(file_obj)
    except Exception as exc:
        raise ParsingError(f"Не удалось открыть PDF-файл: {exc}") from exc

    if len(reader.pages) == 0:
        raise ParsingError("PDF-файл не содержит страниц.")

    pages: list[str] = []
    for page_num, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
            text = text.strip()
            if text:
                pages.append(text)
        except Exception as exc:
            logger.warning("Не удалось извлечь текст со страницы %d: %s", page_num, exc)

    if not pages:
        raise ParsingError(
            "PDF не содержит извлекаемого текста. "
            "Возможно, документ является сканом без текстового слоя."
        )

    return "\n\n".join(pages)


def _parse_docx(file_obj: BinaryIO) -> str:
    """
    Извлекает текст из DOCX-файла.

    Обходит все параграфы документа. Пустые параграфы (разрывы страниц)
    заменяются двойным переносом строки для сохранения структуры.
    """
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise RuntimeError(
            "Библиотека python-docx не установлена. Выполните: poetry add python-docx"
        ) from exc

    try:
        doc = DocxDocument(file_obj)
    except Exception as exc:
        raise ParsingError(f"Не удалось открыть DOCX-файл: {exc}") from exc

    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Дополнительно извлекаем текст из таблиц
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                paragraphs.append(" | ".join(row_cells))

    if not paragraphs:
        raise ParsingError("DOCX-файл не содержит текстового содержимого.")

    return "\n\n".join(paragraphs)


# ---------------------------------------------------------------------------
# Публичный интерфейс
# ---------------------------------------------------------------------------

def extract_text_from_file(file_obj: BinaryIO, extension: str) -> str:
    """
    Основная точка входа: извлекает текст из файла в зависимости от расширения.

    Args:
        file_obj:  открытый бинарный файловый объект (seekable).
        extension: расширение в нижнем регистре, включая точку (напр. ".pdf").

    Returns:
        Строка с извлечённым текстом (гарантированно непустая).

    Raises:
        UnsupportedFileTypeError: расширение не из списка ALLOWED_EXTENSIONS.
        ParsingError:             файл повреждён или не содержит текста.
        RuntimeError:             зависимость не установлена.
    """
    ext = extension.lower()

    if ext not in ALLOWED_EXTENSIONS:
        raise UnsupportedFileTypeError(
            f"Формат «{ext}» не поддерживается. "
            f"Допустимые форматы: {', '.join(sorted(ALLOWED_EXTENSIONS))}."
        )

    logger.debug("Начало парсинга файла с расширением %s", ext)

    if ext == ".pdf":
        text = _parse_pdf(file_obj)
    elif ext == ".docx":
        text = _parse_docx(file_obj)
    else:
        # Недостижимая ветка (защита от будущих изменений ALLOWED_EXTENSIONS)
        raise UnsupportedFileTypeError(f"Обработчик для «{ext}» не реализован.")

    logger.debug("Парсинг завершён: извлечено %d символов", len(text))
    return text
